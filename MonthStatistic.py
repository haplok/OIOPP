import docx
import os
from datetime import date, datetime
import calendar

import MyFunctions as functions



def createMonthStatisticDocx (counterByDayISO, listInOut):
	mainPath = os.getcwd()

	dicIn = counterByDayISO [0]
	dicOut = counterByDayISO [1]

	# print(dicIn["2020-01-20"])

	doc = docx.Document('MonthStatistic.docx')
	l = list(dicIn.elements())
	dateOfStat = date.fromisoformat(l[0])
	
	
	monthS = dateOfStat.month
	yearS = dateOfStat.year

	dic = {}
	month = functions.rusMonthIm (monthS)
	monthYear = f'{month} {yearS} года'
	dic.update({'M_MonthYearOfStatistic':monthYear})

	for element in dic:
		editElement = element.split('_')[1]
		
		for paragraph in doc.paragraphs:
			if str(editElement) in paragraph.text:
				inline = paragraph.runs
				for i in range(len(inline)):
					if str(editElement) in inline[i].text:
						text = inline[i].text.replace(str(editElement), str(dic[element]))
						inline[i].text = text

	table = doc.tables[0]

	

	maxDay = calendar.monthrange(yearS, monthS)[1]
	if dateOfStat.year == date.today().year and dateOfStat.month ==date.today().month:
		maxDay = date.today().day

	countIn = 0
	countOut = 0

	for row in range(maxDay):
		dateVar = date(yearS, monthS, row + 1)
		dateVarISO = str(dateVar.isoformat())
		# print (dateVarISO)
		# dicTable = listDictOut [row]
		table.add_row()
		cell0 = table.cell(row + 2, 0)
		cell0.text = str(row + 1)

		

		cell1 = table.cell(row + 2, 1)
		cell1.text = str(dicIn[dateVarISO])

		cell2 = table.cell(row +2, 2)
		countIn += dicIn[dateVarISO]
		cell2.text = str(countIn)

		cell3 = table.cell(row + 2, 3)
		cell3.text = str(dicOut[dateVarISO])

		cell4 = table.cell(row + 2, 4)
		countOut += dicOut[dateVarISO]
		cell4.text = str(countOut)
		
		pairCounterAverage = counterAveragePatientsByDay (dateVar, listInOut)
		cell5 = table.cell(row + 2, 5) 
		cell5.text = str(pairCounterAverage[0])

		cell6 = table.cell(row + 2, 6)
		cell6.text = str(pairCounterAverage[1])

		# cell1 = tableOut.cell(row + 1, 1)
		# FIO = dicTable['lineEdit_PatientLastName'] + ' ' + dicTable['lineEdit_PatientFirstName'] + ' ' + dicTable['lineEdit_PatientPatronymic']
		# cell1.text = str(FIO)

		# cell2 = tableOut.cell(row + 1, 2)
		# cell2.text = functions.editingDateTime(dicTable['dateEdit_PatientBirthDate'])[0]

		# cell3 = tableOut.cell(row + 1, 3)
		# regAdress = ''
		# if dicTable['lineEdit_RoadReg'] != '':
		# 	regAdress = '\nАдрес прописки: '+ functions.combineAdress(dicTable, 'reg')
		# else:
		# 	pass
		# cell3.text = functions.combineAdress(dicTable, 'main') + regAdress

		# cell4 = tableOut.cell(row + 1, 4)
		# cell4.text = str(dicTable['lineEdit_OurMainDiagnosis'])

		# cell5 = tableOut.cell(row + 1, 5)
		# organization = dicTable['comboBox_MoveToOrganization']
		# department = dicTable['lineEdit_MoveToDepatment']
		# doctor = dicTable['lineEdit_MoveToDoctor']
		# cell5.text = f'{organization}, отделение {department}, врач {doctor}'
					
	
	

	
	nameToSave = f'stat{yearS}g{monthS}'
	
	path = f'documents/{yearS}/month_statistic/'
	
	try:
		os.makedirs(path)
	except:
		pass
	
	doc.save(f'{path}{nameToSave}.docx')
	os.chdir(path)
	os.startfile(f'{nameToSave}.docx')
	os.chdir(mainPath)

def counterAveragePatientsByDay(dateVar, listInOut):
	countOfPatients = [0,0]
	summDaysAll = 0

	if listInOut != None:
		for i in range(len(listInOut)):
			dayIn = functions.dateFromIso(listInOut[i][0])
			dayOut = functions.dateFromIso(listInOut[i][1])
						
			if dayIn.date() <= dateVar < dayOut.date():
				countOfPatients[0] += 1

				summdays = dateVar - dayIn.date()
				summDaysAll += summdays.days + 2
				
			if dayOut.date() <= dayIn.date() <= dateVar:
				countOfPatients[0] += 1

				summdays = dateVar - dayIn.date()
				summDaysAll += summdays.days + 2
				
				# summDaysAll += daysInOIOPPList(listInOut[i])
		if countOfPatients [0] == 0:
			aver = 0
		else:
			aver = summDaysAll / countOfPatients [0]
		countOfPatients[1] = "%.2f" % round(aver, 2)
			# elif dayOut.date() <= dayIn.date():
			# 	countOfPatients += 1

	return countOfPatients

def daysInOIOPPList (tupleInOut):
	inISO = str(functions.dateFromIso(tupleInOut[0]))
	outISO = str(functions.dateFromIso(tupleInOut[1]))


	dtIn = datetime.fromisoformat(inISO)
	dtOut = datetime.fromisoformat(outISO)

	dt = dtOut-dtIn
	if dt.days > 0:
		return (dt.days + 2)
	else:
		return 0
	

