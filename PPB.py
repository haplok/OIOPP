import docx
from openpyxl import Workbook, load_workbook
import os
from datetime import datetime, date

import MyFunctions as functions



def pPBFirstEdition(dictionaryToFill):
	mainPath = os.getcwd()

	doc = docx.Document('PPBFirst.docx')
	dictionaryToFill.pop('id')
	dic = dictionaryToFill

	ppb1Date = functions.dateOfBothPPB(dic)[0]
	ppb2Date = functions.dateOfBothPPB(dic)[1]

	ppb1Date = functions.editDate(ppb1Date)
	ppb2Date = functions.editDate(ppb2Date)

	age = functions.calculationAge(dic['dateEdit_PatientBirthDate'])
	bDate = str(functions.editingDateTime(dic['dateEdit_PatientBirthDate'])[0])
	adressMain = functions.combineAdress(dic, 'main')
	if dic['lineEdit_RoadReg'] == '':
		regAdress = ''
	else:
		regAdress = 'Адрес прописки: '+ functions.combineAdress(dic, 'reg')
	dic.update({'A_AdressMain':adressMain})
	dic.update({'A_AdressReg':regAdress})

	radioButonCheckDic = functions.radioButtonCheck (dic)

	dic.update(radioButonCheckDic)




	dic.update({'A_Age':age})
	dic.update({'dateEdit_PatientBirthDate':bDate})


	dic.update({'P_PPB1Date':ppb1Date,'P_PPB2Date':ppb2Date})
	

	for element in dic:
		editElement = element.split('_')[1]
		
		for paragraph in doc.paragraphs:
			if str(editElement) in paragraph.text:
				inline = paragraph.runs
				for i in range(len(inline)):
					if str(editElement) in inline[i].text:
						text = inline[i].text.replace(str(editElement), str(dic[element]))
						inline[i].text = text
					# print (inline[i].text)

	caseNumber = str(dic['lineEdit_HistoryNumber']).rjust(3,'0')
	fullLastName = str(dic['lineEdit_PatientLastName'])[:1].upper() + str(dic['lineEdit_PatientLastName'])[1:]
	shortFirstName = str(dic['lineEdit_PatientFirstName'])[:1].upper()
	shortPatronymic = str(dic['lineEdit_PatientPatronymic'])[:1].upper()
	nameToSave = f'{caseNumber}_{fullLastName}{shortFirstName}{shortPatronymic}'
	# print (nameToSave)
	yearOfHistory = str(dic['lineEdit_YearOfHistory'])
	path = f'documents/20{yearOfHistory}/ppb1/'
	try:
		os.makedirs(path)
	except:
		pass
	
	doc.save(f'{path}{nameToSave}.docx')
	os.chdir(path)
	os.startfile(f'{nameToSave}.docx')
	os.chdir(mainPath)

def pPBSecondEdition (dictionaryToFill):

	mainPath = os.getcwd()

	doc = docx.Document('PPBSecond.docx')
	dictionaryToFill.pop('id')
	dic = dictionaryToFill

	ppb1Date = functions.dateOfBothPPB(dic)[0]
	ppb2Date = functions.dateOfBothPPB(dic)[1]

	ppb1Date = functions.editDate(ppb1Date)
	ppb2Date = functions.editDate(ppb2Date)

	age = functions.calculationAge(dic['dateEdit_PatientBirthDate'])
	bDate = str(functions.editingDateTime(dic['dateEdit_PatientBirthDate'])[0])
	adressMain = functions.combineAdress(dic, 'main')
	if dic['lineEdit_RoadReg'] == '':
		regAdress = ''
	else:
		regAdress = 'Адрес прописки: '+ functions.combineAdress(dic, 'reg')
	dic.update({'A_AdressMain':adressMain})
	dic.update({'A_AdressReg':regAdress})

	radioButonCheckDic = functions.radioButtonCheck (dic)

	dic.update(radioButonCheckDic)

	dic.update({'A_Age':age})
	dic.update({'dateEdit_PatientBirthDate':bDate})


	dic.update({'P_PPB1Date':ppb1Date,'P_PPB2Date':ppb2Date})

	for element in dic:
		editElement = element.split('_')[1]
		
		for paragraph in doc.paragraphs:
			if str(editElement) in paragraph.text:
				inline = paragraph.runs
				for i in range(len(inline)):
					if str(editElement) in inline[i].text:
						text = inline[i].text.replace(str(editElement), str(dic[element]))
						inline[i].text = text
					# print (inline[i].text)

	caseNumber = str(dic['lineEdit_HistoryNumber']).rjust(3,'0')
	fullLastName = str(dic['lineEdit_PatientLastName'])[:1].upper() + str(dic['lineEdit_PatientLastName'])[1:]
	shortFirstName = str(dic['lineEdit_PatientFirstName'])[:1].upper()
	shortPatronymic = str(dic['lineEdit_PatientPatronymic'])[:1].upper()
	nameToSave = f'{caseNumber}_{fullLastName}{shortFirstName}{shortPatronymic}'
	# print (nameToSave)
	yearOfHistory = str(dic['lineEdit_YearOfHistory'])
	path = f'documents/20{yearOfHistory}/ppb2/'
	try:
		os.makedirs(path)
	except:
		pass
	
	doc.save(f'{path}{nameToSave}.docx')
	os.chdir(path)
	os.startfile(f'{nameToSave}.docx')
	os.chdir(mainPath)
	
def pPBTableEdition (datePPB, firstPPBList, secondPPBList):


	# day = date.fromisoformat(dayISO)


	
	mainPath = os.getcwd()

	loadWB = load_workbook('PPBTabl.xlsx')
	wSheet = loadWB.active


	dateOfPPB = wSheet ['R4']
	dateOfPPB.value = datePPB.date()

	wSheet.merge_cells('E6:T6')
	firstPPB = wSheet ['E6']
	firstPPB.value = 'ПЕРВИЧНЫЕ ППБ:'

	deltaRow = 7
	for i in range (deltaRow, deltaRow + len(firstPPBList)):
		numberCell = wSheet [f'A{i}']
		fioCell = wSheet [f'B{i}']
		bDateCell = wSheet [f'C{i}']
		historyNumbrCell = wSheet [f'D{i}']
		diagnosisCell = wSheet [f'E{i}']
		dateInCell = wSheet [f'F{i}']
		examinationCell = wSheet [f'G{i}']
		conversationCell = wSheet [f'H{i}']
		epiCell = wSheet [f'I{i}']
		familyConsultationCell = wSheet [f'J{i}']
		individualCell = wSheet [f'K{i}']
		psychoStudyCell = wSheet [f'L{i}']
		motivationCell = wSheet [f'M{i}']
		tbnCell = wSheet [f'N{i}']
		cognitCell = wSheet [f'O{i}']
		workCell = wSheet [f'P{i}']
		mobilityCell = wSheet [f'T{i}']
		


		numberCell.value = i - deltaRow + 1
		fioCell.value = f'{firstPPBList [i - deltaRow][2]} \
		{firstPPBList [i - deltaRow][3]} \
		{firstPPBList [i - deltaRow][4]}'
		bDateCell.value = datetime.fromisoformat(firstPPBList [i - deltaRow][7]).date()
		historyNumbrCell.value = f'{firstPPBList [i - deltaRow][0]}/{firstPPBList [i - deltaRow][1]}'
		diagnosisCell.value = firstPPBList [i - deltaRow][8]
		dateInCell.value = datetime.fromisoformat(firstPPBList [i - deltaRow][5]).date()
		examinationCell.value = 1
		conversationCell.value = 1
		epiCell.value = int(firstPPBList [i - deltaRow][9])
		familyConsultationCell.value = int(firstPPBList [i - deltaRow][10])
		individualCell.value = int(firstPPBList [i - deltaRow][11])
		psychoStudyCell.value = int(firstPPBList [i - deltaRow][12])
		motivationCell.value = int(firstPPBList [i - deltaRow][13])
		tbnCell.value = int(firstPPBList [i - deltaRow][14])
		cognitCell.value = 0
		workCell.value = 0
		if firstPPBList [i - deltaRow][15] == '1':
			mobilityCell.value = 'М1'
		elif firstPPBList [i - deltaRow][16] == '1':
			mobilityCell.value = 'М2'
		elif firstPPBList [i - deltaRow][17] == '1':
			mobilityCell.value = 'М3'
		elif firstPPBList [i - deltaRow][18] == '1':
			mobilityCell.value = 'М4'
		elif firstPPBList [i - deltaRow][19] == '1':
			mobilityCell.value = 'Н/М'

	deltaRow += len(firstPPBList)

	wSheet.merge_cells(f'E{deltaRow}:T{deltaRow}')
	secondPPB = wSheet [f'E{deltaRow}']
	secondPPB.value = 'ПОВТОРНЫЕ ППБ:'

	deltaRow += 1

	for i in range (deltaRow, deltaRow + len(secondPPBList)):
		numberCell = wSheet [f'A{i}']
		fioCell = wSheet [f'B{i}']
		bDateCell = wSheet [f'C{i}']
		historyNumbrCell = wSheet [f'D{i}']
		diagnosisCell = wSheet [f'E{i}']
		dateInCell = wSheet [f'F{i}']
		
		mobilityCell = wSheet [f'T{i}']
		socialRoute = wSheet [f'U{i}']
		


		numberCell.value = i - deltaRow + 1
		fioCell.value = f'{secondPPBList [i - deltaRow][2]} \
		{secondPPBList [i - deltaRow][3]} \
		{secondPPBList [i - deltaRow][4]}'
		bDateCell.value = datetime.fromisoformat(secondPPBList [i - deltaRow][7]).date()
		historyNumbrCell.value = f'{secondPPBList [i - deltaRow][0]}/{secondPPBList [i - deltaRow][1]}'
		diagnosisCell.value = secondPPBList [i - deltaRow][8]
		dateInCell.value = datetime.fromisoformat(secondPPBList [i - deltaRow][5]).date()
		
		if secondPPBList [i - deltaRow][15] == '1':
			mobilityCell.value = 'М1'
		elif secondPPBList [i - deltaRow][16] == '1':
			mobilityCell.value = 'М2'
		elif secondPPBList [i - deltaRow][17] == '1':
			mobilityCell.value = 'М3'
		elif secondPPBList [i - deltaRow][18] == '1':
			mobilityCell.value = 'М4'
		elif secondPPBList [i - deltaRow][19] == '1':
			mobilityCell.value = 'Н/М'
		socialRoute.value = secondPPBList [i - deltaRow][20]


	yearOfPPB = datePPB.date().year
	
		
	path = f'documents/{yearOfPPB}/ppbTable/'
	nameToSave = f'ППБ_{str(datePPB)[:10]}'
	try:
		os.makedirs(path)
	except:
		pass
	
	loadWB.save(f'{path}{nameToSave}.xlsx')
	os.chdir(path)
	os.startfile(f'{nameToSave}.xlsx')
	os.chdir(mainPath)


	

	





