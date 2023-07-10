import docx
import os
import random
from datetime import datetime
from datetime import timedelta

import MyFunctions as functions




def coverCaseHistory(dictionaryToFill):

	mainPath = os.getcwd()

	doc = docx.Document('b-003u.docx')
	dictionaryToFill.pop('id')
	dic = dictionaryToFill

	dateInISO = dic['dateTimeEdit_MoveFromDateTime']
	daysIn = functions.daysInOIOPP(dictionaryToFill)
	age = functions.calculationAge(dic['dateEdit_PatientBirthDate'])
	dateTimeMoveFrom = str(functions.editingDateTime(dic['dateTimeEdit_MoveFromDateTime'])[0])+\
	 '   ' + str(functions.editingDateTime(dic['dateTimeEdit_MoveFromDateTime'])[1])
	dateMoveFrom = str(functions.editingDateTime(dic['dateTimeEdit_MoveFromDateTime'])[0])	
	if daysIn == '':
		dateMoveTo = ''
	else:
		dateMoveTo = str(functions.editingDateTime(dic['dateEdit_MoveToDate'])[0])
	bDate = str(functions.editingDateTime(dic['dateEdit_PatientBirthDate'])[0])
	dic.update({'dateEdit_PatientBirthDate':bDate,
				'dateTimeEdit_MoveFromDateTime':dateTimeMoveFrom, 
				'dateEdit_MoveToDate':dateMoveTo}) 
	dic.update({'lineEdit_DaysInOIOPP':daysIn})
	dic.update({'D_OnlyDateMoveFrom':dateMoveFrom})

	mobility = ''
	if dic['radioButton_GroupM1'] == '1':
		mobility = 'М1'
	elif dic['radioButton_GroupM2'] == '1':
		mobility = 'М2'
	elif dic['radioButton_GroupM3'] == '1':
		mobility = 'М3'
	elif dic['radioButton_GroupImmobile'] == '1':
		mobility = 'Немобильный пациент'
	dic.update({'M_Mobility':mobility})

	disabilityGroup = ''
	if dic['radioButton_DisabilityGroup0'] == '1':
		disabilityGroup = 'Нет группы инвалидности'
	elif dic['radioButton_DisabilityGroup1'] == '1':
		disabilityGroup = 'I группа инвалидности'
	elif dic['radioButton_DisabilityGroup2'] == '1':
		disabilityGroup = 'II группа инвалидности'
	elif dic['radioButton_DisabilityGroup3'] == '1':
		disabilityGroup = 'III группа инвалидности'

	if dic['dateEdit_FLGDate'] == '2000-01-01T00:00:00':
		dic['dateEdit_FLGDate'] = '__.__.20__г.'
	else: dic['dateEdit_FLGDate'] = str(functions.editingDateTime(dic['dateEdit_FLGDate'])[0])


	dic.update({'D_DisabilityGroup':disabilityGroup})
	

	FIO = dic['lineEdit_PatientLastName'] + ' ' + dic['lineEdit_PatientFirstName'] \
	+ ' ' + dic['lineEdit_PatientPatronymic']
	dic.update({'F_FIO':FIO.title()})

	sex = ''
	if dic['radioButton_Male'] == '1':
		sex = 'Муж'
	elif dic['radioButton_Female'] == '1':
		sex = 'Жен'

	adressMain = functions.combineAdress(dic, 'main')
		
	entrFlorTel = entranceFloorTelephone(
		dic['lineEdit_EntranceMain'],dic['lineEdit_FloorMain'],
		dic['lineEdit_CodeMain'],dic['lineEdit_HomeTelephoneNumberMain'],
		dic['lineEdit_MobilePhoneNumberMain'])
	mainAdressFull = adressMain + entrFlorTel
	
	pso = functions.passportSnilsOmsGen(dic['lineEdit_PassportID'], 
		dic['lineEdit_SnilsID'], dic['lineEdit_OmsID']) 
	dic['lineEdit_PassportID'] = pso[0]
	dic['lineEdit_SnilsID'] = pso[1]
	dic['lineEdit_OmsID'] = pso[2]

	if dic['lineEdit_RoadReg'] == '':
		regAdress = ''
	else:
		regAdress = 'Адрес прописки: '+ functions.combineAdress(dic, 'reg')
			

	fullProxy = proxyGenerator (
		dic['comboBox_ProxyType'], dic['lineEdit_ProxyLastName'],
		dic['lineEdit_ProxyFirstName'], dic['lineEdit_ProxyPatronymic'],
		dic['lineEdit_ProxyTelephoneNumber1'], dic['lineEdit_ProxyTelephoneNumber2'],
		dic['lineEdit_ProxyTelephoneNumber3'])
	dic.update({'A_AdressMain':mainAdressFull})
	dic.update({'A_AdressShortMain':adressMain})
	dic.update({'A_AdressReg':regAdress})
	dic.update({'S_Sex':sex})
	dic.update({'F_FullProxy':fullProxy})

	dic.update({'A_Age':age})

	temperature = str(random.randrange(0,10,1)/10 + 36.0)

	dic.update({'T_temperature':temperature})

	
	for element in dic:
		editElement = element.split('_')[1]
		
		for paragraph in doc.paragraphs:
			if str(editElement) in paragraph.text:
				inline = paragraph.runs
				for i in range(len(inline)):
					if str(editElement) in inline[i].text:
						text = inline[i].text.replace(str(editElement), str(dic[element]))
						inline[i].text = text
					# print (inline[i].text)

	tableListNazn1 = doc.tables[4]
	tableListNazn2 = doc.tables[5]
	
	nextVar = None

	for col in range (2, len(tableListNazn1.columns)):
		monthDay = daysMonthGen (dateInISO, col - 2)

		cell0 = tableListNazn1.cell(0, col)
		cell0.text = monthDay[0]
		prevcell = tableListNazn1.cell(0, col - 1)
		
		if cell0.text == prevcell.text:
			prevcell.merge(cell0)
			prevcell.text = monthDay[0]

		cell1 = tableListNazn1.cell(1, col)
		cell1.text = monthDay[1]
		nextVar = col

	for col in range (2, len(tableListNazn2.columns)):
		monthDay = daysMonthGen (dateInISO, col + nextVar - 3)

		cell0 = tableListNazn2.cell(0, col)
		cell0.text = monthDay[0]
		prevcell = tableListNazn2.cell(0, col - 1)
		
		if cell0.text == prevcell.text:
			prevcell.merge(cell0)
			prevcell.text = monthDay[0]

		cell1 = tableListNazn2.cell(1, col)
		cell1.text = monthDay[1]
		
	# tableListNazn2.cell(5,5).getCellFormat().getShading().setBackgroundPatternColor(Color.RED)
		

	

	caseNumber = str(dic['lineEdit_HistoryNumber']).rjust(3,'0')
	fullLastName = str(dic['lineEdit_PatientLastName'])[:1].upper() + str(dic['lineEdit_PatientLastName'])[1:]
	shortFirstName = str(dic['lineEdit_PatientFirstName'])[:1].upper()
	shortPatronymic = str(dic['lineEdit_PatientPatronymic'])[:1].upper()
	nameToSave = f'{caseNumber}_{fullLastName}{shortFirstName}{shortPatronymic}'
	# print (nameToSave)
	yearOfHistory = str(dic['lineEdit_YearOfHistory'])
	path = f'documents/20{yearOfHistory}/covers/'
	try:
		os.makedirs(path)
	except:
		pass
	
	doc.save(f'{path}{nameToSave}.docx')
	os.chdir(path)
	os.startfile(f'{nameToSave}.docx')
	os.chdir(mainPath)

def daysMonthGen (dateISO, dayDelta):
	baseDate = datetime.fromisoformat(dateISO)
	currDate = baseDate + timedelta (days = dayDelta)

	month = functions.rusMonthIm(currDate.month)
	day = str(currDate.day)

	editMonthDay = (month, day)

	return editMonthDay




def entranceFloorTelephone(entrance, floor, code, homeTelephone, mobTelephone):
	if homeTelephone != '':
		homeTelephone = f'дом. телефон {homeTelephone}\t\t'
	if mobTelephone != '':
		mobTelephone = f'сот. телефон {mobTelephone}\t\t'
	efc = f' подъезд {entrance}, этаж {floor}, домофон {code}\n{homeTelephone}{mobTelephone}'
	return efc

def proxyGenerator(proxy, lastName, firstName, patronymic, tel1, tel2, tel3):

	fullProxy = f"""{proxy} {lastName} {firstName} {patronymic}\t{tel1}\t{tel2}\t{tel3}"""

	return fullProxy



	


