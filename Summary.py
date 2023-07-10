

import docx
import os
from datetime import datetime

import MyFunctions as functions

def createDaySummary(listDictIn, listDictOut, dateOfSummaryISO):

	mainPath = os.getcwd()
	doc = docx.Document('Summary.docx')
	# dictionaryToFill.pop('id')
	dic = {}
	dateRusSum = functions.editingDateTime(dateOfSummaryISO)[0]
	dic.update({'D_DateOfSummary':dateRusSum})

	for element in dic:
		editElement = element.split('_')[1]
		
		for paragraph in doc.paragraphs:
			if str(editElement) in paragraph.text:
				inline = paragraph.runs
				for i in range(len(inline)):
					if str(editElement) in inline[i].text:
						text = inline[i].text.replace(str(editElement), str(dic[element]))
						inline[i].text = text
					# print (inline[i].text)

	tableIn = doc.tables[0]

	for row in range(len(listDictIn)):
		dicTable = listDictIn [row]
		tableIn.add_row()
		cell0 = tableIn.cell(row + 1, 0)
		cell0.text = str(row + 1)

		cell1 = tableIn.cell(row + 1, 1)
		FIO = dicTable['lineEdit_PatientLastName'] + ' ' + dicTable['lineEdit_PatientFirstName'] \
	+ ' ' + dicTable['lineEdit_PatientPatronymic']
		cell1.text = str(FIO)

		cell2 = tableIn.cell(row + 1, 2)
		cell2.text = functions.editingDateTime(dicTable['dateEdit_PatientBirthDate'])[0]

		cell3 = tableIn.cell(row + 1, 3)
		regAdress = ''
		if dicTable['lineEdit_RoadReg'] != '':
			regAdress = '\nАдрес прописки: '+ functions.combineAdress(dicTable, 'reg')
		else:
			pass
		cell3.text = functions.combineAdress(dicTable, 'main') + regAdress

		cell4 = tableIn.cell(row + 1, 4)
		cell4.text = str(dicTable['lineEdit_MainDiagnosisAnotherDepartment'])

		cell5 = tableIn.cell(row + 1, 5)
		organization = dicTable['comboBox_MoveFromOrganization']
		department = dicTable['lineEdit_MoveFromDepartment']
		doctor = dicTable['lineEdit_MoveFromDoctor']
		cell5.text = f'{organization}, отделение {department}, врач {doctor}'

	tableOut = doc.tables[1]

	for row in range(len(listDictOut)):
		dicTable = listDictOut [row]
		tableOut.add_row()
		cell0 = tableOut.cell(row + 1, 0)
		cell0.text = str(row + 1)

		cell1 = tableOut.cell(row + 1, 1)
		FIO = dicTable['lineEdit_PatientLastName'] + ' ' + dicTable['lineEdit_PatientFirstName'] + ' ' + dicTable['lineEdit_PatientPatronymic']
		cell1.text = str(FIO)

		cell2 = tableOut.cell(row + 1, 2)
		cell2.text = functions.editingDateTime(dicTable['dateEdit_PatientBirthDate'])[0]

		cell3 = tableOut.cell(row + 1, 3)
		regAdress = ''
		if dicTable['lineEdit_RoadReg'] != '':
			regAdress = '\nАдрес прописки: '+ functions.combineAdress(dicTable, 'reg')
		else:
			pass
		cell3.text = functions.combineAdress(dicTable, 'main') + regAdress

		cell4 = tableOut.cell(row + 1, 4)
		cell4.text = str(dicTable['lineEdit_OurMainDiagnosis'])

		cell5 = tableOut.cell(row + 1, 5)
		organization = dicTable['comboBox_MoveToOrganization']
		department = dicTable['lineEdit_MoveToDepatment']
		doctor = dicTable['lineEdit_MoveToDoctor']
		cell5.text = f'{organization}, отделение {department}, врач {doctor}'





	dateOfSummary = datetime.fromisoformat(dateOfSummaryISO)

	dayS = dateOfSummary.day
	monthS = dateOfSummary.month
	yearS = dateOfSummary.year

	
	nameToSave = f'Сводка{yearS}г{monthS}м{dayS}'
	
	path = f'documents/{yearS}/Summary/'
	try:
		os.makedirs(path)
	except:
		pass
	
	doc.save(f'{path}{nameToSave}.docx')
	os.chdir(path)
	os.startfile(f'{nameToSave}.docx')
	os.chdir(mainPath)